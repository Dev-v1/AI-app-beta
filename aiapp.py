import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk, filedialog
import requests
import json
import os
from datetime import datetime
import threading
import time
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def markdown_to_plain_text(md: str) -> str:
    if not md:
        return ""

    text = md

    # remove fenced code block markers but keep code content
    text = re.sub(r"```[^\n]*\n", "", text)  # opening fence with optional language
    text = text.replace("```", "")

    # inline code: `code` -> code
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # bold and italic markers
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*([^*\n]+)\*", r"\1", text)    # *italic*
    text = re.sub(r"__([^_]+)__", r"\1", text)      # __bold__
    text = re.sub(r"_([^_\n]+)_", r"\1", text)      # _italic_

    # headings: remove leading #'s
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", text)

    # blockquotes: remove leading >
    text = re.sub(r"(?m)^\s{0,3}>\s?", "", text)

    # list markers: keep text, remove bullets and numbers
    text = re.sub(r"(?m)^\s*[-*+]\s+", "", text)
    text = re.sub(r"(?m)^\s*\d+\.\s+", "", text)

    # links: [text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)

    # images: ![alt](url) -> alt
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1", text)

    # horizontal rules
    text = re.sub(r"(?m)^\s*(-{3,}|\*{3,}|_{3,})\s*$", "", text)

    # collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# Optional dependency: tl_markdown_converter
# If it isn't importable in the current interpreter, fall back to a minimal built-in converter.
try:
    from tl_markdown_converter import MarkdownConverter  # type: ignore
except ModuleNotFoundError:
    class MarkdownConverter:
        """
        Minimal fallback converter.

        Produces a list of segments (dict) with keys:
          text, bold, italic, code, code_block, header, list_item

        Not a full Markdown implementation, just enough for basic chat formatting.
        """

        _INLINE_PATTERN = re.compile(
            r"(?P<code>`[^`]+`)"
            r"|(?P<bold>\*\*[^*]+\*\*)"
            r"|(?P<italic>\*[^*\n]+\*)"
        )

        def convert(self, text: str) -> list[dict]:
            segments: list[dict] = []
            if not text:
                return segments

            lines = text.splitlines(keepends=True)
            in_code_block = False

            for line in lines:
                stripped = line.strip()

                # Toggle fenced code blocks
                if stripped.startswith("```"):
                    in_code_block = not in_code_block
                    continue

                if in_code_block:
                    segments.append({"text": line, "code_block": True})
                    continue

                # Headers
                if stripped.startswith("#"):
                    segments.append({"text": line, "header": True})
                    continue

                # Lists
                if stripped.startswith(("- ", "* ", "+ ")) or re.match(r"^\d+\.\s", stripped):
                    segments.append({"text": line, "list_item": True})
                    continue

                # Inline formatting
                segments.extend(self._convert_inline(line))

            return segments

        def _convert_inline(self, line: str) -> list[dict]:
            out: list[dict] = []
            pos = 0

            for m in self._INLINE_PATTERN.finditer(line):
                start, end = m.span()
                if start > pos:
                    out.append({"text": line[pos:start]})

                token = line[start:end]
                if m.group("code"):
                    out.append({"text": token[1:-1], "code": True})
                elif m.group("bold"):
                    out.append({"text": token[2:-2], "bold": True})
                elif m.group("italic"):
                    out.append({"text": token[1:-1], "italic": True})

                pos = end

            if pos < len(line):
                out.append({"text": line[pos:]})

            return out


# Configuration files
CONFIG_FILE = "chat_config.json"
CHATS_FILE = "chat_history.json"


class OpenAIChat:
    def __init__(self, api_key=None, server_url="https://api.openai.com/v1", model="gpt-3.5-turbo"):
        self.api_key = api_key
        self.server_url = server_url.rstrip('/')
        self.model = model
        self.conversation_history = []

    def send_message(self, message, max_retries=3):
        self.conversation_history.append({"role": "user", "content": message})

        headers = {"Content-Type": "application/json"}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"

        data = {"model": self.model, "messages": self.conversation_history}

        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f"{self.server_url}/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=30
                )
                response.raise_for_status()

                result = response.json()
                assistant_message = result['choices'][0]['message']['content']
                self.conversation_history.append({"role": "assistant", "content": assistant_message})
                return assistant_message

            except requests.exceptions.HTTPError as e:
                if getattr(e.response, "status_code", None) == 429:
                    if attempt < max_retries - 1:
                        wait_time = (2 ** attempt) * 2
                        time.sleep(wait_time)
                        continue
                    return "Error: Rate limit exceeded. Please wait a moment and try again. (429 Too Many Requests)"
                return f"Error: {str(e)}"
            except Exception as e:
                return f"Error: {str(e)}"

        return "Error: Maximum retry attempts reached"

    def clear_history(self):
        self.conversation_history = []

    def get_models(self):
        headers = {"Content-Type": "application/json"}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"

        try:
            response = requests.get(f"{self.server_url}/models", headers=headers, timeout=30)
            response.raise_for_status()
            result = response.json()
            return [model['id'] for model in result.get('data', [])]
        except Exception:
            return []


OPENROUTER_POPULAR_MODELS = [
    "anthropic/claude-sonnet-4",
    "anthropic/claude-3.5-haiku",
    "openai/gpt-4o",
    "openai/gpt-4o-mini",
    "google/gemini-2.0-flash-001",
    "google/gemini-2.5-pro-preview",
    "meta-llama/llama-3.3-70b-instruct",
    "deepseek/deepseek-chat",
    "mistralai/mistral-large",
]


class ChatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Chat Assistant")
        self.root.geometry("1200x800")

        self.bg_color = "#1e1e1e"
        self.sidebar_color = "#252525"
        self.input_bg = "#2d2d2d"
        self.text_color = "#ffffff"
        self.accent_color = "#007acc"
        self.button_color = "#0e639c"

        self.root.configure(bg=self.bg_color)

        self.chat_sessions = {}
        self.current_chat_id = None

        main_container = tk.Frame(root, bg=self.bg_color)
        main_container.pack(fill=tk.BOTH, expand=True)

        sidebar = tk.Frame(main_container, bg=self.sidebar_color, width=280, padx=10, pady=10)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)

        title_label = tk.Label(
            sidebar, text="AI Chat Assistant", font=("Arial", 16, "bold"),
            bg=self.sidebar_color, fg=self.text_color
        )
        title_label.pack(pady=(0, 20))

        self.new_chat_btn = tk.Button(
            sidebar, text="+ New Chat", command=self.create_new_chat,
            bg=self.accent_color, fg=self.text_color, font=("Arial", 11, "bold"),
            relief=tk.FLAT, cursor="hand2", padx=20, pady=10
        )
        self.new_chat_btn.pack(fill=tk.X, pady=(0, 20))

        chat_list_label = tk.Label(
            sidebar, text="Chat History", font=("Arial", 11, "bold"),
            bg=self.sidebar_color, fg=self.text_color
        )
        chat_list_label.pack(anchor=tk.W, pady=(0, 5))

        chat_list_container = tk.Frame(sidebar, bg=self.sidebar_color)
        chat_list_container.pack(fill=tk.BOTH, expand=True, pady=(0, 20))

        self.chat_listbox = tk.Listbox(
            chat_list_container, bg=self.input_bg, fg=self.text_color,
            font=("Arial", 10), relief=tk.FLAT, selectbackground=self.accent_color,
            highlightthickness=0, activestyle='none'
        )
        self.chat_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.chat_listbox.bind('<<ListboxSelect>>', self.switch_chat)

        chat_scrollbar = tk.Scrollbar(chat_list_container, command=self.chat_listbox.yview)
        chat_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.chat_listbox.config(yscrollcommand=chat_scrollbar.set)

        self.delete_chat_btn = tk.Button(
            sidebar, text="Delete Chat", command=self.delete_chat,
            bg="#5c2d2d", fg=self.text_color, font=("Arial", 10),
            relief=tk.FLAT, cursor="hand2", padx=10, pady=8
        )
        self.delete_chat_btn.pack(fill=tk.X, pady=(0, 10))

        self.export_word_btn = tk.Button(
            sidebar, text="Export to Word", command=self.export_to_word,
            bg="#2d5c2d", fg=self.text_color, font=("Arial", 10),
            relief=tk.FLAT, cursor="hand2", padx=10, pady=8
        )
        self.export_word_btn.pack(fill=tk.X, pady=(0, 20))

        config_frame = tk.Frame(sidebar, bg=self.sidebar_color)
        config_frame.pack(fill=tk.X, pady=(0, 10))

        config_label = tk.Label(
            config_frame, text="Configuration", font=("Arial", 11, "bold"),
            bg=self.sidebar_color, fg=self.text_color
        )
        config_label.pack(anchor=tk.W, pady=(0, 10))

        self.mode_var = tk.StringVar(value="openrouter")
        mode_label = tk.Label(config_frame, text="Mode:", bg=self.sidebar_color, fg=self.text_color, font=("Arial", 9))
        mode_label.pack(anchor=tk.W, pady=(0, 5))

        mode_btn_frame = tk.Frame(config_frame, bg=self.sidebar_color)
        mode_btn_frame.pack(fill=tk.X, pady=(0, 10))
        tk.Radiobutton(
            mode_btn_frame, text="OpenRouter", variable=self.mode_var, value="openrouter",
            command=self.toggle_mode, bg=self.sidebar_color, fg=self.text_color,
            selectcolor=self.input_bg, font=("Arial", 9)
        ).pack(anchor=tk.W)
        tk.Radiobutton(
            mode_btn_frame, text="Remote API", variable=self.mode_var, value="remote",
            command=self.toggle_mode, bg=self.sidebar_color, fg=self.text_color,
            selectcolor=self.input_bg, font=("Arial", 9)
        ).pack(anchor=tk.W)
        tk.Radiobutton(
            mode_btn_frame, text="Local Model", variable=self.mode_var, value="local",
            command=self.toggle_mode, bg=self.sidebar_color, fg=self.text_color,
            selectcolor=self.input_bg, font=("Arial", 9)
        ).pack(anchor=tk.W)

        self.server_url_label = tk.Label(config_frame, text="Server URL:", bg=self.sidebar_color, fg=self.text_color, font=("Arial", 9))
        self.server_url_entry = tk.Entry(config_frame, bg=self.input_bg, fg=self.text_color, font=("Arial", 9), relief=tk.FLAT, insertbackground=self.text_color)
        self.server_url_entry.insert(0, "https://api.openai.com/v1")

        self.api_key_label = tk.Label(config_frame, text="API Key:", bg=self.sidebar_color, fg=self.text_color, font=("Arial", 9))
        self.api_key_entry = tk.Entry(config_frame, show="*", bg=self.input_bg, fg=self.text_color, font=("Arial", 9), relief=tk.FLAT, insertbackground=self.text_color)

        self.local_model_label = tk.Label(config_frame, text="Local Model:", bg=self.sidebar_color, fg=self.text_color, font=("Arial", 9))
        self.local_model_entry = tk.Entry(config_frame, bg=self.input_bg, fg=self.text_color, font=("Arial", 9), relief=tk.FLAT, insertbackground=self.text_color)
        self.local_model_button = tk.Button(
            config_frame, text="Browse", command=self.browse_local_model,
            bg=self.button_color, fg=self.text_color, font=("Arial", 8),
            relief=tk.FLAT, cursor="hand2"
        )

        self.model_label = tk.Label(config_frame, text="Model:", bg=self.sidebar_color, fg=self.text_color, font=("Arial", 9))
        self.model_var = tk.StringVar()

        self.model_combobox = ttk.Combobox(config_frame, textvariable=self.model_var, font=("Arial", 9))
        self.model_combobox['values'] = OPENROUTER_POPULAR_MODELS

        self.model_dropdown = tk.OptionMenu(config_frame, self.model_var, "")
        self.model_dropdown.config(state=tk.DISABLED)

        self.connect_button = tk.Button(
            config_frame, text="Connect & Go", command=self.connect,
            bg=self.accent_color, fg=self.text_color, font=("Arial", 10, "bold"),
            relief=tk.FLAT, cursor="hand2", padx=10, pady=8
        )

        self.fetch_models_button = tk.Button(
            config_frame, text="Fetch Models",
            command=self.fetch_openrouter_models,
            bg=self.button_color, fg=self.text_color, font=("Arial", 8),
            relief=tk.FLAT, cursor="hand2"
        )

        chat_container = tk.Frame(main_container, bg=self.bg_color)
        chat_container.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        try:
            import platform
            if platform.system() == "Windows":
                emoji_font = ("Segoe UI Emoji", 11)
            elif platform.system() == "Darwin":
                emoji_font = ("Apple Color Emoji", 11)
            else:
                emoji_font = ("Noto Color Emoji", 11)
        except Exception:
            emoji_font = ("Arial", 11)

        self.chat_display = scrolledtext.ScrolledText(
            chat_container, wrap=tk.WORD, state=tk.DISABLED,
            bg=self.input_bg, fg=self.text_color,
            font=emoji_font, relief=tk.FLAT,
            insertbackground=self.text_color, padx=10, pady=10
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        try:
            import platform
            if platform.system() == "Windows":
                user_font = ("Segoe UI Emoji", 11, "bold")
                assistant_font = ("Segoe UI Emoji", 11, "bold")
                system_font = ("Segoe UI Emoji", 10, "italic")
                timestamp_font = ("Segoe UI Emoji", 8)
                thinking_font = ("Segoe UI Emoji", 10, "italic")
            elif platform.system() == "Darwin":
                user_font = ("Apple Color Emoji", 11, "bold")
                assistant_font = ("Apple Color Emoji", 11, "bold")
                system_font = ("Apple Color Emoji", 10, "italic")
                timestamp_font = ("Apple Color Emoji", 8)
                thinking_font = ("Apple Color Emoji", 10, "italic")
            else:
                user_font = ("Noto Color Emoji", 11, "bold")
                assistant_font = ("Noto Color Emoji", 11, "bold")
                system_font = ("Noto Color Emoji", 10, "italic")
                timestamp_font = ("Noto Color Emoji", 8)
                thinking_font = ("Noto Color Emoji", 10, "italic")
        except Exception:
            user_font = ("Arial", 11, "bold")
            assistant_font = ("Arial", 11, "bold")
            system_font = ("Arial", 10, "italic")
            timestamp_font = ("Arial", 8)
            thinking_font = ("Arial", 10, "italic")

        self.chat_display.tag_config("user", foreground="#7dd3fc", font=user_font)
        self.chat_display.tag_config("assistant", foreground="#86efac", font=assistant_font)
        self.chat_display.tag_config("system", foreground="#fbbf24", font=system_font)
        self.chat_display.tag_config("timestamp", foreground="#94a3b8", font=timestamp_font)
        self.chat_display.tag_config("thinking", foreground="#a78bfa", font=thinking_font)

        self.chat_display.tag_config("bold", font=(emoji_font[0], emoji_font[1], "bold"))
        self.chat_display.tag_config("italic", font=(emoji_font[0], emoji_font[1], "italic"))
        self.chat_display.tag_config("code", foreground="#f472b6", background="#374151", font=("Courier", 10))
        self.chat_display.tag_config("code_block", foreground="#e5e7eb", background="#1f2937", font=("Courier", 10))
        self.chat_display.tag_config("header", font=(emoji_font[0], emoji_font[1] + 2, "bold"), foreground="#60a5fa")
        self.chat_display.tag_config("list_item", lmargin1=20, lmargin2=20)

        self.thinking_active = False
        self.thinking_mark = None

        input_frame = tk.Frame(chat_container, bg=self.bg_color)
        input_frame.pack(fill=tk.X)

        self.message_entry = tk.Entry(
            input_frame, bg=self.input_bg, fg=self.text_color,
            font=("Arial", 11), relief=tk.FLAT, insertbackground=self.text_color
        )
        self.message_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10), ipady=8)
        self.message_entry.bind("<Return>", lambda e: self.send_message())
        self.message_entry.config(state=tk.DISABLED)

        self.send_button = tk.Button(
            input_frame, text="Send", command=self.send_message, state=tk.DISABLED,
            bg=self.accent_color, fg=self.text_color, font=("Arial", 10, "bold"),
            relief=tk.FLAT, cursor="hand2", padx=20, pady=8
        )
        self.send_button.pack(side=tk.LEFT, padx=(0, 5))

        self.clear_button = tk.Button(
            input_frame, text="Clear", command=self.clear_chat, state=tk.DISABLED,
            bg="#5c2d2d", fg=self.text_color, font=("Arial", 10),
            relief=tk.FLAT, cursor="hand2", padx=15, pady=8
        )
        self.clear_button.pack(side=tk.LEFT)

        self.chat_client = None

        self.load_config()
        self.load_chats()
        self.update_chat_list()
        self.toggle_mode()

    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    config = json.load(f)
                self.api_key_entry.delete(0, tk.END)
                self.api_key_entry.insert(0, config.get('api_key', ''))
                self.server_url_entry.delete(0, tk.END)
                self.server_url_entry.insert(0, config.get('server_url', 'https://api.openai.com/v1'))
                self.mode_var.set(config.get('mode', 'openrouter'))
                model = config.get('model', '')
                if model:
                    self.model_var.set(model)
            except Exception as e:
                print(f"Error loading config: {e}")

    def save_config(self):
        config = {
            'api_key': self.api_key_entry.get().strip(),
            'server_url': self.server_url_entry.get().strip(),
            'mode': self.mode_var.get(),
            'model': self.model_var.get()
        }
        try:
            with open(CONFIG_FILE, 'w') as f:
                json.dump(config, f, indent=2)
        except Exception as e:
            print(f"Error saving config: {e}")

    def load_chats(self):
        if os.path.exists(CHATS_FILE):
            try:
                with open(CHATS_FILE, 'r') as f:
                    self.chat_sessions = json.load(f)
            except Exception as e:
                print(f"Error loading chats: {e}")
                self.chat_sessions = {}

    def save_chats(self):
        try:
            with open(CHATS_FILE, 'w') as f:
                json.dump(self.chat_sessions, f, indent=2)
        except Exception as e:
            print(f"Error saving chats: {e}")

    def create_new_chat(self):
        chat_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        chat_title = f"Chat {datetime.now().strftime('%b %d, %H:%M')}"

        self.chat_sessions[chat_id] = {
            'title': chat_title,
            'messages': [],
            'created': datetime.now().isoformat()
        }

        self.current_chat_id = chat_id
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.delete(1.0, tk.END)
        self.chat_display.config(state=tk.DISABLED)

        if self.chat_client:
            self.chat_client.clear_history()

        self.save_chats()
        self.update_chat_list()

        for i, (cid, _) in enumerate(sorted(self.chat_sessions.items(), key=lambda x: x[1]['created'], reverse=True)):
            if cid == chat_id:
                self.chat_listbox.selection_clear(0, tk.END)
                self.chat_listbox.selection_set(i)
                break

    def switch_chat(self, event=None):
        selection = self.chat_listbox.curselection()
        if not selection:
            return

        index = selection[0]
        chat_ids = sorted(self.chat_sessions.keys(), key=lambda x: self.chat_sessions[x]['created'], reverse=True)

        if index < len(chat_ids):
            self.current_chat_id = chat_ids[index]

            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)

            chat_data = self.chat_sessions[self.current_chat_id]
            for msg in chat_data['messages']:
                self.chat_display.insert(tk.END, msg)

            self.chat_display.config(state=tk.DISABLED)
            self.chat_display.see(tk.END)

            if self.chat_client:
                self.chat_client.clear_history()
                for msg in chat_data['messages']:
                    if msg.startswith("You: "):
                        content = msg.replace("You: ", "").split("\n")[0]
                        self.chat_client.conversation_history.append({"role": "user", "content": content})
                    elif msg.startswith("Assistant: "):
                        content = msg.replace("Assistant: ", "").split("\n")[0]
                        self.chat_client.conversation_history.append({"role": "assistant", "content": content})

    def delete_chat(self):
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("Warning", "Please select a chat to delete")
            return

        if messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this chat?"):
            index = selection[0]
            chat_ids = sorted(self.chat_sessions.keys(), key=lambda x: x[1]['created'], reverse=True)

            if index < len(chat_ids):
                chat_id = chat_ids[index]
                del self.chat_sessions[chat_id]

                if self.current_chat_id == chat_id:
                    self.current_chat_id = None
                    self.chat_display.config(state=tk.NORMAL)
                    self.chat_display.delete(1.0, tk.END)
                    self.chat_display.config(state=tk.DISABLED)
                    if self.chat_client:
                        self.chat_client.clear_history()

                self.save_chats()
                self.update_chat_list()

    def update_chat_list(self):
        self.chat_listbox.delete(0, tk.END)
        sorted_chats = sorted(self.chat_sessions.items(), key=lambda x: x[1]['created'], reverse=True)
        for _, chat_data in sorted_chats:
            self.chat_listbox.insert(tk.END, chat_data['title'])

    def toggle_mode(self):
        mode = self.mode_var.get()

        self.server_url_label.pack_forget()
        self.server_url_entry.pack_forget()
        self.api_key_label.pack_forget()
        self.api_key_entry.pack_forget()
        self.local_model_label.pack_forget()
        self.local_model_entry.pack_forget()
        self.local_model_button.pack_forget()
        self.model_label.pack_forget()
        self.model_dropdown.pack_forget()
        self.model_combobox.pack_forget()
        self.fetch_models_button.pack_forget()
        self.connect_button.pack_forget()

        if mode == "openrouter":
            self.api_key_label.pack(anchor=tk.W, pady=(0, 2))
            self.api_key_entry.pack(fill=tk.X, pady=(0, 10), ipady=4)
            self.model_label.pack(anchor=tk.W, pady=(0, 2))
            self.model_combobox.pack(fill=tk.X, pady=(0, 5), ipady=4)
            self.fetch_models_button.pack(fill=tk.X, pady=(0, 10))
            self.connect_button.pack(fill=tk.X)
            if not self.model_var.get():
                self.model_var.set(OPENROUTER_POPULAR_MODELS[0])

        elif mode == "remote":
            self.server_url_label.pack(anchor=tk.W, pady=(0, 2))
            self.server_url_entry.pack(fill=tk.X, pady=(0, 10), ipady=4)
            self.api_key_label.pack(anchor=tk.W, pady=(0, 2))
            self.api_key_entry.pack(fill=tk.X, pady=(0, 10), ipady=4)
            self.model_label.pack(anchor=tk.W, pady=(0, 2))
            self.model_dropdown.pack(fill=tk.X, pady=(0, 10))
            self.connect_button.pack(fill=tk.X)

        else:
            self.local_model_label.pack(anchor=tk.W, pady=(0, 2))
            model_entry_frame = tk.Frame(self.local_model_entry.master, bg=self.sidebar_color)
            model_entry_frame.pack(fill=tk.X, pady=(0, 10))
            self.local_model_entry.pack(in_=model_entry_frame, side=tk.LEFT, fill=tk.X, expand=True, ipady=4)
            self.local_model_button.pack(in_=model_entry_frame, side=tk.RIGHT, padx=(5, 0))
            self.model_label.pack(anchor=tk.W, pady=(0, 2))
            self.model_dropdown.pack(fill=tk.X, pady=(0, 10))
            self.connect_button.pack(fill=tk.X)

    def browse_local_model(self):
        filename = filedialog.askopenfilename(title="Select Local Model", filetypes=(("All files", "*.*"),))
        if filename:
            self.local_model_entry.delete(0, tk.END)
            self.local_model_entry.insert(0, filename)

    def fetch_openrouter_models(self):
        api_key = self.api_key_entry.get().strip()
        temp_client = OpenAIChat(api_key if api_key else None, "https://openrouter.ai/api/v1")
        models = temp_client.get_models()

        if models:
            models.sort()
            self.model_combobox['values'] = models
            messagebox.showinfo("Success", f"Fetched {len(models)} models from OpenRouter")
        else:
            messagebox.showwarning("Warning", "Could not fetch models. Using popular models list.")
            self.model_combobox['values'] = OPENROUTER_POPULAR_MODELS

    def connect(self):
        mode = self.mode_var.get()

        if mode == "openrouter":
            api_key = self.api_key_entry.get().strip()
            if not api_key:
                messagebox.showerror("Error", "Please enter your OpenRouter API key")
                return

            selected_model = self.model_var.get().strip()
            if not selected_model:
                messagebox.showerror("Error", "Please select or enter a model")
                return

            server_url = "https://openrouter.ai/api/v1"
            self.chat_client = OpenAIChat(api_key, server_url, selected_model)

            if not self.current_chat_id:
                self.create_new_chat()

            self.append_to_chat(f"System: Connected to OpenRouter with model {selected_model}\n\n")

        elif mode == "remote":
            server_url = self.server_url_entry.get().strip()
            api_key = self.api_key_entry.get().strip()

            if not server_url:
                messagebox.showerror("Error", "Please enter a server URL")
                return

            temp_client = OpenAIChat(api_key if api_key else None, server_url)
            models = temp_client.get_models()

            if not models:
                messagebox.showerror("Error", "Could not fetch models from server")
                return

            menu = self.model_dropdown["menu"]
            menu.delete(0, "end")
            for model in models:
                menu.add_command(label=model, command=lambda m=model: self.model_var.set(m))

            self.model_var.set(models[0])
            self.model_dropdown.config(state=tk.NORMAL)

            selected_model = self.model_var.get()
            self.chat_client = OpenAIChat(api_key if api_key else None, server_url, selected_model)

            if not self.current_chat_id:
                self.create_new_chat()

            self.append_to_chat(f"System: Connected to server with model {selected_model}\n\n")

        else:
            local_model_path = self.local_model_entry.get().strip()

            if not local_model_path:
                messagebox.showerror("Error", "Please select a local model file")
                return

            if not os.path.exists(local_model_path):
                messagebox.showerror("Error", "Local model file does not exist")
                return

            server_url = "http://localhost:8080/v1"
            model_name = os.path.basename(local_model_path)

            self.chat_client = OpenAIChat(None, server_url, model_name)
            self.model_dropdown.config(state=tk.DISABLED)

            if not self.current_chat_id:
                self.create_new_chat()

            self.append_to_chat(f"System: Connected to local model {model_name}\n\n")

        self.message_entry.config(state=tk.NORMAL)
        self.send_button.config(state=tk.NORMAL)
        self.clear_button.config(state=tk.NORMAL)

        self.save_config()
        messagebox.showinfo("Success", "Connected successfully!")

    def send_message(self):
        if not self.chat_client:
            messagebox.showerror("Error", "Please connect to a server first")
            return

        if not self.current_chat_id:
            self.create_new_chat()

        message = self.message_entry.get().strip()
        if not message:
            return

        is_first_message = len([msg for msg in self.chat_sessions[self.current_chat_id]['messages'] if msg.startswith("You: ")]) == 0

        self.append_to_chat(f"You: {message}\n\n")
        self.message_entry.delete(0, tk.END)

        self.send_button.config(state=tk.DISABLED)
        self.root.update()

        self.show_thinking()

        def get_response():
            response = self.chat_client.send_message(message)

            def apply_ui_updates():
                self.hide_thinking()
                self.append_to_chat(f"Assistant: {response}\n\n")
                if is_first_message:
                    self.auto_name_chat(message)
                self.send_button.config(state=tk.NORMAL)
                self.save_chats()
                self.update_chat_list()

            self.root.after(0, apply_ui_updates)

        threading.Thread(target=get_response, daemon=True).start()

    def clear_chat(self):
        if not self.current_chat_id:
            return

        if messagebox.askyesno("Confirm Clear", "Are you sure you want to clear this chat?"):
            if self.chat_client:
                self.chat_client.clear_history()

            self.chat_sessions[self.current_chat_id]['messages'] = []

            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            self.chat_display.config(state=tk.DISABLED)

            self.append_to_chat("System: Chat cleared\n\n")
            self.save_chats()

    def auto_name_chat(self, first_message):
        try:
            naming_client = OpenAIChat(
                api_key=self.chat_client.api_key,
                server_url=self.chat_client.server_url,
                model=self.chat_client.model
            )

            naming_prompt = (
                f"Generate a very short title (3-5 words max) for a chat that starts with: "
                f"'{first_message[:100]}'. Respond with ONLY the title, nothing else."
            )
            chat_name = naming_client.send_message(naming_prompt).strip()
            chat_name = chat_name.strip('"\'').strip()
            if len(chat_name) > 50:
                chat_name = chat_name[:47] + "..."

            if self.current_chat_id and self.current_chat_id in self.chat_sessions:
                self.chat_sessions[self.current_chat_id]['title'] = chat_name

        except Exception as e:
            print(f"Auto naming failed: {e}")

    def show_thinking(self):
        self.thinking_active = True
        self.chat_display.config(state=tk.NORMAL)
        self.thinking_mark = self.chat_display.index(tk.END + "-1c")
        self.chat_display.insert(tk.END, "Assistant is thinking", "thinking")
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.see(tk.END)

        def animate():
            patterns = [".", "..", "...", "....", "...", ".."]
            state = 0

            while self.thinking_active:
                def step():
                    if not self.thinking_active or not self.thinking_mark:
                        return
                    self.chat_display.config(state=tk.NORMAL)
                    self.chat_display.delete(self.thinking_mark, f"{self.thinking_mark} lineend")
                    self.chat_display.insert(
                        self.thinking_mark,
                        f"Assistant is thinking{patterns[state % len(patterns)]}",
                        "thinking"
                    )
                    self.chat_display.config(state=tk.DISABLED)
                    self.chat_display.see(tk.END)

                self.root.after(0, step)
                state += 1
                time.sleep(0.1)

        threading.Thread(target=animate, daemon=True).start()

    def hide_thinking(self):
        self.thinking_active = False
        if self.thinking_mark:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(self.thinking_mark, f"{self.thinking_mark} lineend")
            self.chat_display.config(state=tk.DISABLED)
            self.thinking_mark = None

    def render_markdown(self, text):
        converter = MarkdownConverter()
        formatted_segments = converter.convert(text)

        for segment in formatted_segments:
            content = segment.get('text', '')
            tags = []

            if segment.get('bold'):
                tags.append('bold')
            if segment.get('italic'):
                tags.append('italic')
            if segment.get('code'):
                tags.append('code')
            if segment.get('code_block'):
                tags.append('code_block')
            if segment.get('header'):
                tags.append('header')
            if segment.get('list_item'):
                tags.append('list_item')

            if tags:
                self.chat_display.insert(tk.END, content, tuple(tags) if len(tags) > 1 else tags[0])
            else:
                self.chat_display.insert(tk.END, content)

    def render_inline_markdown(self, text):
        converter = MarkdownConverter()
        formatted_segments = converter.convert(text)

        for segment in formatted_segments:
            content = segment.get('text', '')
            tags = []

            if segment.get('bold'):
                tags.append('bold')
            if segment.get('italic'):
                tags.append('italic')
            if segment.get('code'):
                tags.append('code')

            if tags:
                self.chat_display.insert(tk.END, content, tuple(tags) if len(tags) > 1 else tags[0])
            else:
                self.chat_display.insert(tk.END, content)

    def append_to_chat(self, text):
        self.chat_display.config(state=tk.NORMAL)

        timestamp = datetime.now().strftime("%H:%M")

        if text.startswith("You: "):
            self.chat_display.insert(tk.END, "You", "user")
            self.chat_display.insert(tk.END, f" [{timestamp}]", "timestamp")
            self.chat_display.insert(tk.END, text[4:])

        elif text.startswith("Assistant: "):
            self.chat_display.insert(tk.END, "Assistant", "assistant")
            self.chat_display.insert(tk.END, f" [{timestamp}]", "timestamp")
            self.render_markdown(text[11:])

        elif text.startswith("System: "):
            self.chat_display.insert(tk.END, text, "system")

        else:
            self.chat_display.insert(tk.END, text)

        self.chat_display.see(tk.END)
        self.chat_display.config(state=tk.DISABLED)

        if self.current_chat_id and self.current_chat_id in self.chat_sessions:
            self.chat_sessions[self.current_chat_id]['messages'].append(text)

    def export_to_word(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror(
                "Error",
                "python-docx library is not installed.\n\nInstall it with: pip install python-docx"
            )
            return

        if not self.current_chat_id:
            messagebox.showwarning("Warning", "No chat selected to export")
            return

        chat_data = self.chat_sessions[self.current_chat_id]
        if not chat_data['messages']:
            messagebox.showwarning("Warning", "Chat is empty")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            initialfile=f"{chat_data['title']}.docx"
        )

        if not filename:
            return

        try:
            doc = Document()

            # Add title
            title = doc.add_heading(chat_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            created_date = datetime.fromisoformat(chat_data['created']).strftime("%B %d, %Y at %H:%M")
            metadata = doc.add_paragraph(f"Created: {created_date}")
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metadata_format = metadata.runs[0].font
            metadata_format.size = Pt(10)
            metadata_format.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()

            for message in chat_data['messages']:
                if message.startswith("You: "):
                    content = message[5:].strip()
                    p = doc.add_paragraph()
                    user_run = p.add_run("You: ")
                    user_run.bold = True
                    user_run.font.color.rgb = RGBColor(125, 211, 252)
                    user_run.font.size = Pt(11)

                    content_run = p.add_run(content)
                    content_run.font.size = Pt(11)

                elif message.startswith("Assistant: "):
                    raw_content = message[11:].strip()
                    content = markdown_to_plain_text(raw_content)

                    p = doc.add_paragraph()
                    assistant_run = p.add_run("Assistant: ")
                    assistant_run.bold = True
                    assistant_run.font.color.rgb = RGBColor(134, 239, 172)
                    assistant_run.font.size = Pt(11)

                    content_run = p.add_run(content)
                    content_run.font.size = Pt(11)

                elif message.startswith("System: "):
                    content = message[8:].strip()
                    p = doc.add_paragraph()
                    system_run = p.add_run(f"System: {content}")
                    system_run.italic = True
                    system_run.font.color.rgb = RGBColor(251, 191, 36)
                    system_run.font.size = Pt(10)

                doc.add_paragraph()

            doc.save(filename)
            messagebox.showinfo("Success", f"Chat exported successfully to:\n{filename}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to export chat:\n{str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = ChatApp(root)
    root.mainloop()